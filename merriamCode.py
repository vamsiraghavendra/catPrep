import json
import csv
import os
import re
import time
from urllib.parse import unquote
from pathlib import Path
from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout
from playwright.sync_api import Error as PWError
import requests
try:
    from docx import Document
    from docx.shared import Inches
except Exception:  # pragma: no cover
    Document = None
    Inches = None

def _get_secret(name: str) -> str:
    value = os.getenv(name, "").strip()
    if value:
        return value
    try:
        import streamlit as st

        return str(st.secrets.get(name, "")).strip()
    except Exception:
        return ""

SAVED_WORDS_URL = "https://www.merriam-webster.com/saved-words"
LOGIN_URL = "https://www.merriam-webster.com/login"

OUT_DIR = Path("D:/catSfotwareDev/merriam2Story")
OUT_JSON = OUT_DIR / "words.json"
OUT_CSV = OUT_DIR / "words.csv"
OUT_DOCX = OUT_DIR / "words.docx"

MW_THESAURUS_KEY = _get_secret("MW_THESAURUS_KEY")
MW_DICTIONARY_KEY = _get_secret("MW_DICTIONARY_KEY")

DICT_ENDPOINT = "https://dictionaryapi.com/api/v3/references/collegiate/json/{word}"
THES_ENDPOINT = "https://dictionaryapi.com/api/v3/references/thesaurus/json/{word}"
REQUEST_TIMEOUT_S = 20
SLEEP_BETWEEN_WORDS_S = 0.15
MAX_DEFS_PER_WORD = 5
MAX_SYNS_PER_WORD = 25
MAX_EXAMPLES_PER_WORD = 2

# Set to True to stop on the Saved Words page so you can visually confirm your list loads.
# Set to False to continue with scraping/export.
WAIT_ONLY = False

# Set True to save debug artifacts if extraction returns 0 (HTML + screenshot + console samples)
DEBUG_DUMP_ON_EMPTY = True

# Set True to print per-page diagnostics about what selectors matched.
DEBUG_LOG = True


def unique_preserve_order(items):
    seen = set()
    out = []
    for x in items:
        x = (x or "").strip()
        if not x:
            continue
        k = x.lower()
        if k in seen:
            continue
        seen.add(k)
        out.append(x)
    return out


def _is_suggestions_payload(payload) -> bool:
    # MW returns a list; when no entries, it's often a list of strings (suggestions)
    return isinstance(payload, list) and (len(payload) == 0 or isinstance(payload[0], str))


def clean_mw_text(text: str) -> str:
    if not text:
        return ""
    # Replace {bc} (bold colon) with standard colon and space
    text = text.replace("{bc}", ": ")
    # Clean tags like {sx|blend||} -> blend
    text = re.sub(r"\{sx\|(.*?)\|.*?\}", r"\1", text)
    # Clean links {d_link|text|target} -> text
    text = re.sub(r"\{(?:d_link|a_link)\|([^|}]*).*?\}", r"\1", text)
    # Remove simple formatting tags like {it}, {wi}, {b}, {inf}, {sup}, {sc}
    text = re.sub(r"\{/?(?:wi|it|b|inf|sup|sc|parahw)\}", "", text)
    # Remove any remaining tags
    text = re.sub(r"\{.*?\}", "", text)
    # Collapse multiple spaces
    text = re.sub(r"\s+", " ", text).strip()
    # Remove leading colon/space if present (common after {bc} removal)
    if text.startswith(": "):
        text = text[2:]
    # Capitalize first letter
    if text:
        text = text[0].upper() + text[1:]
    return text


def fetch_word_details(word: str, session: requests.Session) -> list[dict]:
    if not MW_DICTIONARY_KEY:
        raise RuntimeError("Missing `MW_DICTIONARY_KEY` secret.")
    url = DICT_ENDPOINT.format(word=word)
    resp = session.get(url, params={"key": MW_DICTIONARY_KEY}, timeout=REQUEST_TIMEOUT_S)
    resp.raise_for_status()
    payload = resp.json()
    if _is_suggestions_payload(payload):
        return []

    senses = []
    for entry in payload:
        def_section = entry.get("def") or []
        for sseq in def_section:
            sseq_list = sseq.get("sseq") or []
            for grouping in sseq_list:
                # Each grouping contains sense items or pseq items
                for item in grouping:
                    if not isinstance(item, list) or len(item) < 2:
                        continue
                    type_label = item[0]
                    data = item[1]

                    # We want to extract data from 'sense', 'bs' (binding sense), or inside 'pseq'
                    targets = []
                    if type_label in ("sense", "bs"):
                        targets.append(data)
                    elif type_label == "pseq":
                        for p_item in data:
                            if isinstance(p_item, list) and len(p_item) > 1 and p_item[0] in ("sense", "bs"):
                                targets.append(p_item[1])

                    for sense_data in targets:
                        dt = sense_data.get("dt") or []
                        
                        # dt is a list of ["type", "content"] blocks
                        def_text_parts = []
                        vis_entries = []
                        for dt_item in dt:
                            if not isinstance(dt_item, list) or len(dt_item) < 2:
                                continue
                            if dt_item[0] == "text":
                                def_text_parts.append(dt_item[1])
                            elif dt_item[0] == "vis":
                                vis_entries.extend(dt_item[1])

                        clean_def = clean_mw_text("".join(def_text_parts))
                        if not clean_def:
                            continue

                        examples = []
                        for vis in vis_entries:
                            t = vis.get("t")
                            if t:
                                ex_text = clean_mw_text(t)
                                auth = vis.get("aq", {}).get("auth", "")
                                examples.append({"t": ex_text, "auth": auth})

                        senses.append({
                            "sn": sense_data.get("sn", "").strip(),
                            "def": clean_def,
                            "examples": examples
                        })
    # Respect the limit for definitions per word
    return senses[:MAX_DEFS_PER_WORD]


def fetch_synonyms(word: str, session: requests.Session) -> list[str]:
    if not MW_THESAURUS_KEY:
        raise RuntimeError("Missing `MW_THESAURUS_KEY` secret.")
    url = THES_ENDPOINT.format(word=word)
    resp = session.get(url, params={"key": MW_THESAURUS_KEY}, timeout=REQUEST_TIMEOUT_S)
    resp.raise_for_status()
    payload = resp.json()
    if _is_suggestions_payload(payload):
        return []

    syns: list[str] = []
    for entry in payload:
        meta = entry.get("meta") or {}
        syn_groups = meta.get("syns")
        # syns is typically list[list[str]]
        if isinstance(syn_groups, list):
            for grp in syn_groups:
                if isinstance(grp, list):
                    syns.extend([str(x).strip() for x in grp if str(x).strip()])
        if len(syns) >= MAX_SYNS_PER_WORD:
            break
    return unique_preserve_order(syns)[:MAX_SYNS_PER_WORD]


def write_docx(words: list[str], out_path: Path):
    if Document is None:
        raise RuntimeError("python-docx is not installed. Run: python -m pip install python-docx")
    doc = Document()
    doc.add_heading("Merriam-Webster Saved Words", level=1)
    doc.add_paragraph(f"Generated from {len(words)} words scraped from Saved Words.")

    with requests.Session() as session:
        for i, word in enumerate(words, start=1):
            doc.add_heading(f"{i}. {word}", level=2)

            try:
                details = fetch_word_details(word, session)
            except Exception as e:
                details = []
                doc.add_paragraph(f"Definition lookup failed: {e}")

            try:
                syns = fetch_synonyms(word, session)
            except Exception as e:
                syns = []
                doc.add_paragraph(f"Synonym lookup failed: {e}")

            if details:
                for sense in details:
                    # Format: 1a. Definition text
                    label = sense['sn']
                    if label and not label.endswith("."):
                        label += "."
                    
                    p = doc.add_paragraph()
                    if label:
                        run = p.add_run(f"{label} ")
                        run.bold = True
                    p.add_run(sense['def'])
                    
                    # Examples Indented
                    for ex in sense['examples']:
                        txt = f"\"{ex['t']}\""
                        if ex['auth']:
                            txt += f" — {ex['auth']}"
                        p_ex = doc.add_paragraph(txt)
                        if Inches:
                            p_ex.paragraph_format.left_indent = Inches(0.2)
            else:
                doc.add_paragraph("(no definitions found)")

            doc.add_paragraph("Synonyms:")
            if syns:
                doc.add_paragraph(", ".join(syns))
            else:
                doc.add_paragraph("(no synonyms found)")

            time.sleep(SLEEP_BETWEEN_WORDS_S)

    doc.save(out_path)


WORD_RE = re.compile(r"[A-Za-z][A-Za-z'\- ]{0,60}$")
WORD_LOOSE_RE = re.compile(r"[A-Za-z][A-Za-z'\- ]{0,80}")

# Saved Words can include both Dictionary and Thesaurus entries.
SAVED_WORD_ANCHOR_SEL = "a[href^='/dictionary/'], a[href^='/thesaurus/']"


def _word_from_entry_href(href: str) -> str:
    href = (href or "").strip()
    if not href:
        return ""

    # Handle both relative and absolute URLs; support dictionary + thesaurus.
    # Examples:
    #   /dictionary/seldom
    #   /thesaurus/adjudication
    if "/dictionary/" in href:
        tail = href.split("/dictionary/", 1)[1]
    elif "/thesaurus/" in href:
        tail = href.split("/thesaurus/", 1)[1]
    else:
        return ""

    tail = tail.split("?", 1)[0].split("#", 1)[0]
    tail = tail.strip("/").strip()
    tail = unquote(tail)

    # Some links can contain extra path segments; take the first segment as the word slug
    if "/" in tail:
        tail = tail.split("/", 1)[0]

    return tail.strip()


def extract_words_from_page(page):
    # Saved words are typically anchor text linking to /dictionary/<word>, but the visible text
    # may include extra UI/whitespace. Prefer parsing href when possible.
    words = []
    debug = {"total_anchors": 0, "kept": 0, "sample": []}
    try:
        rows = page.eval_on_selector_all(
            SAVED_WORD_ANCHOR_SEL,
            """els => els.map(e => ({
                t: (e.textContent||'').trim(),
                h: e.getAttribute('href')||''
            }))""",
        )
        debug["total_anchors"] = len(rows or [])
        for it in rows:
            t = (it.get("t") or "").strip()
            h = it.get("h") or ""

            # Prefer href-derived slug
            slug = _word_from_entry_href(h)
            if slug and WORD_RE.fullmatch(slug):
                words.append(slug)
                if len(debug["sample"]) < 8:
                    debug["sample"].append({"from": "href", "word": slug, "href": h, "text": t[:60]})
                continue

            # Fallback: sometimes textContent is clean
            if t and WORD_RE.fullmatch(t):
                words.append(t)
                if len(debug["sample"]) < 8:
                    debug["sample"].append({"from": "text", "word": t, "href": h, "text": t[:60]})
    except Exception:
        pass

    out = unique_preserve_order(words)
    debug["kept"] = len(out)
    return out, debug


def ensure_saved_words_loaded(page, max_scrolls: int = 12):
    """
    Some UIs lazy-load or virtualize lists; scroll to trigger rendering.
    """
    # Wait a bit for initial XHR/render
    try:
        page.wait_for_load_state("networkidle", timeout=8000)
    except Exception:
        pass
    page.wait_for_timeout(800)

    last_count = -1
    for _ in range(max_scrolls):
        try:
            # Scroll down
            page.mouse.wheel(0, 1200)
        except Exception:
            try:
                page.evaluate("() => window.scrollBy(0, 1200)")
            except Exception:
                pass
        page.wait_for_timeout(500)
        try:
            page.wait_for_load_state("networkidle", timeout=4000)
        except Exception:
            pass

        try:
            c = page.locator(SAVED_WORD_ANCHOR_SEL).count()
        except Exception:
            c = -1
        if c == last_count:
            break
        last_count = c


def get_page_counter_text(page) -> str:
    # The UI shows something like "1 of 6"
    for sel in [
        "text=/\\b\\d+\\s+of\\s+\\d+\\b/i",
        ".pagination:text-matches(\"\\\\b\\\\d+\\\\s+of\\\\s+\\\\d+\\\\b\", \"i\")",
    ]:
        try:
            loc = page.locator(sel).first
            if loc.count() > 0:
                t = loc.inner_text().strip()
                if t:
                    return t
        except Exception:
            continue
    return ""


def click_next_if_possible(page) -> bool:
    # Try several “next” selectors; the page shows a » control.
    next_candidates = [
        # Merriam Saved Words pagination often uses <li class="ul-page-next"> as the clickable control.
        "li.ul-page-next",
        "li.ul-page-next[role='button']",
        "a[rel='next']",
        "button[aria-label*='Next' i]",
        "a[aria-label*='Next' i]",
        "button:has-text('»')",
        "a:has-text('»')",
        "button:has-text('›')",
        "a:has-text('›')",
        # Merriam often renders pagination arrows as SVG icons (no visible '»' text).
        "button:has(svg .ul-action-arrow-right)",
        "a:has(svg .ul-action-arrow-right)",
        "button:has(g.ul-action-arrow-right)",
        "a:has(g.ul-action-arrow-right)",
        # Fallback: click the closest button/a ancestor of the icon itself.
        "g.ul-action-arrow-right >> xpath=ancestor::button[1]",
        "g.ul-action-arrow-right >> xpath=ancestor::a[1]",
        "svg .ul-action-arrow-right >> xpath=ancestor::button[1]",
        "svg .ul-action-arrow-right >> xpath=ancestor::a[1]",
    ]

    for sel in next_candidates:
        loc = page.locator(sel).first
        try:
            if loc.count() == 0:
                continue

            # Skip if disabled
            try:
                if loc.is_disabled():
                    continue
            except Exception:
                pass

            # Some sites use aria-disabled
            try:
                aria_disabled = loc.get_attribute("aria-disabled")
                if aria_disabled and aria_disabled.lower() == "true":
                    continue
            except Exception:
                pass

            before_counter = get_page_counter_text(page)
            before_url = page.url
            try:
                before_first = page.locator(SAVED_WORD_ANCHOR_SEL).first.inner_text(timeout=1500).strip()
            except Exception:
                before_first = ""
            loc.scroll_into_view_if_needed()
            try:
                loc.click(timeout=2500)
            except Exception:
                # Some UI layers intercept clicks; force as a last resort.
                loc.click(timeout=2500, force=True)

            # Wait for either the counter to change or the content to settle
            page.wait_for_load_state("domcontentloaded")
            try:
                page.wait_for_load_state("networkidle", timeout=5000)
            except Exception:
                pass
            page.wait_for_timeout(800)

            after_counter = get_page_counter_text(page)
            after_url = page.url
            if before_counter and after_counter and before_counter != after_counter:
                return True

            # Fallback: if URL changed, assume it paged.
            if before_url != after_url:
                return True

            # Fallback: if the first item changes, assume it paged (common for XHR pagination).
            try:
                after_first = page.locator(SAVED_WORD_ANCHOR_SEL).first.inner_text(timeout=2000).strip()
            except Exception:
                after_first = ""
            if before_first and after_first and before_first != after_first:
                return True

            # If neither counter nor URL changed, treat as no pagination.
            return False
        except Exception:
            continue

    # Final fallback: programmatically click the nearest clickable ancestor of the arrow icon.
    try:
        before_counter = get_page_counter_text(page)
        before_url = page.url
        try:
            before_first = page.locator(SAVED_WORD_ANCHOR_SEL).first.inner_text(timeout=1500).strip()
        except Exception:
            before_first = ""

        clicked = page.evaluate(
            """() => {
              const isVisible = (el) => {
                const r = el.getBoundingClientRect();
                const s = window.getComputedStyle(el);
                return !!(r.width && r.height) && s.display !== 'none' && s.visibility !== 'hidden';
              };
              const isDisabled = (el) => {
                if (!el) return true;
                if (el.disabled) return true;
                const ad = (el.getAttribute('aria-disabled') || '').toLowerCase();
                if (ad === 'true') return true;
                return false;
              };

              // Preferred: click the pagination <li class="ul-page-next"> if present.
              const liNext = document.querySelector('li.ul-page-next');
              if (liNext && isVisible(liNext) && !isDisabled(liNext)) {
                liNext.click();
                return true;
              }

              const icons = Array.from(document.querySelectorAll('g.ul-action-arrow-right, .ul-action-arrow-right'));
              for (const icon of icons) {
                if (!isVisible(icon)) continue;
                const clickable = icon.closest('button, a, [role="button"]');
                if (!clickable) continue;
                if (!isVisible(clickable)) continue;
                if (isDisabled(clickable)) continue;
                clickable.click();
                return true;
              }
              return false;
            }"""
        )
        if not clicked:
            return False

        page.wait_for_load_state("domcontentloaded")
        try:
            page.wait_for_load_state("networkidle", timeout=5000)
        except Exception:
            pass
        page.wait_for_timeout(800)

        after_counter = get_page_counter_text(page)
        after_url = page.url
        if before_counter and after_counter and before_counter != after_counter:
            return True
        if before_url != after_url:
            return True
        try:
            after_first = page.locator(SAVED_WORD_ANCHOR_SEL).first.inner_text(timeout=2000).strip()
        except Exception:
            after_first = ""
        if before_first and after_first and before_first != after_first:
            return True
        return False
    except Exception:
        return False


def dismiss_common_overlays(page):
    # Best-effort: cookie/consent banners can block clicks/typing.
    for sel in [
        "button:has-text('Accept')",
        "button:has-text('I Accept')",
        "button:has-text('Agree')",
        "button:has-text('OK')",
        "button#onetrust-accept-btn-handler",
        "button[aria-label*='accept' i]",
        "button[aria-label*='agree' i]",
        "text=Accept All",
        "text=I Agree",
    ]:
        try:
            loc = page.locator(sel).first
            if loc.count() and loc.is_visible():
                loc.click(timeout=1200)
                page.wait_for_timeout(300)
        except Exception:
            pass


def fill_verified(page, selector: str, value: str, label: str):
    loc = page.locator(selector).first
    loc.wait_for(state="visible", timeout=8000)
    loc.scroll_into_view_if_needed()
    try:
        loc.click(timeout=1500)
    except Exception:
        pass

    # Attempt 1: fill
    try:
        loc.fill(value, timeout=3000)
    except Exception:
        pass

    # Verify; if not correct, fallback to typing.
    try:
        current = (loc.input_value(timeout=1500) or "").strip()
    except Exception:
        current = ""

    if current != value:
        try:
            loc.click(timeout=1500)
        except Exception:
            pass
        try:
            loc.press("Control+A")
        except Exception:
            pass
        loc.type(value, delay=35)

    # Final verification (best effort)
    try:
        current2 = (loc.input_value(timeout=1500) or "").strip()
        if current2 != value:
            raise RuntimeError(f"{label} did not populate correctly (got: {current2!r})")
    except PWError:
        # Some inputs might not support input_value; ignore.
        pass


def pick_and_tag_login_identifier_input(page) -> str:
    """
    Tries to find the *actual* visible username/email input by inspecting all inputs.
    Tags the chosen element with data-pw-login-id="1" and returns a CSS selector for it.
    """
    info = page.evaluate(
        """() => {
          const isVisible = (el) => {
            const r = el.getBoundingClientRect();
            const style = window.getComputedStyle(el);
            return !!(r.width && r.height) && style.visibility !== 'hidden' && style.display !== 'none';
          };

          const closestContainer = (el) => {
            if (!el) return null;
            // Prefer the actual form; otherwise use a nearby semantic container.
            return el.closest('form') || el.closest('[role="dialog"]') || el.closest('section') || el.closest('main') || el.parentElement;
          };

          // Clear any previous tag
          for (const el of document.querySelectorAll('[data-pw-login-id="1"]')) {
            el.removeAttribute('data-pw-login-id');
          }

          const passwordEl = document.querySelector('input[type="password"]');
          const pwdContainer = closestContainer(passwordEl);

          const inputs = Array.from(document.querySelectorAll('input'));
          const candidates = inputs.filter(el => {
            if (!isVisible(el)) return false;
            const type = (el.getAttribute('type') || 'text').toLowerCase();
            if (type === 'password') return false;
            if (el.disabled) return false;
            if (el.readOnly) return false;
            // Ignore common non-identifier types
            if (['hidden','submit','button','checkbox','radio','file','search','tel','number','date'].includes(type)) return false;

            // Strongly prefer fields in the same form/container as the password input.
            // This avoids newsletter / footer subscribe fields elsewhere on the page.
            if (pwdContainer) {
              const elContainer = closestContainer(el);
              if (elContainer !== pwdContainer) return false;
            }
            return true;
          });

          const scoreEl = (el) => {
            const s = (v) => (v || '').toString().toLowerCase();
            const type = s(el.getAttribute('type') || 'text');
            const name = s(el.getAttribute('name'));
            const id = s(el.getAttribute('id'));
            const ph = s(el.getAttribute('placeholder'));
            const aria = s(el.getAttribute('aria-label'));
            const ac = s(el.getAttribute('autocomplete'));
            const joined = [type,name,id,ph,aria,ac].join(' ');

            let score = 0;
            if (type === 'email') score += 50;
            if (joined.includes('email')) score += 40;
            if (joined.includes('user')) score += 25;
            if (joined.includes('username')) score += 25;
            if (joined.includes('login')) score += 10;
            if (joined.includes('identifier')) score += 10;
            if (joined.includes('signin')) score += 10;
            // Prefer typical login autocomplete hints
            if (ac.includes('username')) score += 20;
            if (ac.includes('email')) score += 20;
            // Penalize likely irrelevant fields
            if (joined.includes('search')) score -= 50;
            // Hard-exclude marketing/newsletter/subscribe inputs (common culprit).
            if (
              joined.includes('newsletter') ||
              joined.includes('subscribe') ||
              joined.includes('subscription') ||
              joined.includes('marketing') ||
              joined.includes('offers') ||
              joined.includes('promo')
            ) {
              score -= 200;
            }
            return { score, joined, type, name, id, ph, aria, ac };
          };

          const scored = candidates.map(el => ({ el, meta: scoreEl(el) }))
            .sort((a,b) => b.meta.score - a.meta.score);

          const chosen = scored[0];
          if (!chosen) return { ok: false, reason: 'No visible non-password input candidates found', count: candidates.length };

          chosen.el.setAttribute('data-pw-login-id', '1');
          return {
            ok: true,
            count: candidates.length,
            chosen: chosen.meta,
          };
        }"""
    )

    if not info or not info.get("ok"):
        raise RuntimeError(f"Could not auto-pick login identifier input: {info}")

    chosen = info.get("chosen") or {}
    print("Picked login identifier input:", chosen)
    return "input[data-pw-login-id='1']"


def build_word_entries_from_mw(words: list[str]) -> list[dict]:
    """
    For each word, fetch definitions (Dictionary) and synonyms (Thesaurus).
    Returns list of dicts: {word, definition, example_usage, synonyms}
    Same format as aeon_vocab_lookup entries for consistent card display.
    """
    entries = []
    with requests.Session() as session:
        for word in words:
            word = (word or "").strip()
            if not word:
                continue
            definition = ""
            examples: list[str] = []
            synonyms: list[str] = []

            try:
                details = fetch_word_details(word, session)
                if details:
                    definition = details[0].get("def", "")
                    for sense in details[:2]:
                        for ex in sense.get("examples", [])[:MAX_EXAMPLES_PER_WORD]:
                            t = ex.get("t", "")
                            if t:
                                examples.append(t)
            except Exception:
                pass

            try:
                synonyms = fetch_synonyms(word, session)
            except Exception:
                pass

            if not examples and definition:
                examples = [f"Definition: {definition}"]

            entries.append({
                "word": word,
                "definition": definition,
                "example_usage": examples[:2],
                "synonyms": synonyms[:12],
            })
            time.sleep(SLEEP_BETWEEN_WORDS_S)

    return entries


def fetch_saved_words_from_mw(
    email: str | None = None,
    password: str | None = None,
    headless: bool = True,
) -> tuple[list[str], str | None]:
    """
    Log into Merriam-Webster, scrape Saved Words, and return the list.
    Returns (words, error_message). error_message is None on success.
    Email/password must be provided by the caller.
    """
    user_email = (email or "").strip()
    user_password = (password or "").strip()
    if not user_email or not user_password:
        return [], "Email and password are required."

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=headless)
            context = browser.new_context()
            page = context.new_page()

            # Login
            page.goto(LOGIN_URL, wait_until="domcontentloaded")
            dismiss_common_overlays(page)

            filled = False
            for user_sel, pass_sel in [
                ("input#ul-email", "input#ul-password"),
                ("input#ul-email", "input[name='up']"),
                ("input[name='ue']", "input#ul-password"),
                ("input[name='ue']", "input[name='up']"),
                ("input[type='email']", "input[type='password']"),
                ("input[name='email']", "input[name='password']"),
                ("input[name='username']", "input[name='password']"),
                ("input#email", "input#password"),
                ("input#username", "input#password"),
            ]:
                try:
                    page.wait_for_selector(user_sel, timeout=6000, state="visible")
                    page.wait_for_selector(pass_sel, timeout=6000, state="visible")
                    fill_verified(page, user_sel, user_email, "Email/Username")
                    fill_verified(page, pass_sel, user_password, "Password")
                    filled = True
                    break
                except PWTimeout:
                    continue
                except Exception:
                    continue

            if not filled:
                try:
                    page.wait_for_selector("input[type='password']", timeout=8000, state="visible")
                    user_sel = pick_and_tag_login_identifier_input(page)
                    fill_verified(page, user_sel, user_email, "Email/Username")
                    fill_verified(page, "input[type='password']", user_password, "Password")
                    filled = True
                except Exception as e:
                    context.close()
                    browser.close()
                    return [], str(e)

            if not filled:
                context.close()
                browser.close()
                return [], "Could not find login inputs."

            page.keyboard.press("Enter")
            page.wait_for_timeout(1200)
            try:
                page.locator("button[type='submit'], input[type='submit']").first.click(timeout=1500)
            except Exception:
                pass

            page.wait_for_timeout(2000)
            dismiss_common_overlays(page)

            page.goto(SAVED_WORDS_URL, wait_until="domcontentloaded")
            if "login" in page.url.lower():
                context.close()
                browser.close()
                return [], "Login did not stick (redirected to login). Check credentials."

            try:
                page.wait_for_selector(f"{SAVED_WORD_ANCHOR_SEL}, text=/saved\\s+words/i", timeout=15000)
            except Exception:
                pass
            ensure_saved_words_loaded(page)

            all_words = []
            max_pages = 200
            for _ in range(max_pages):
                page.wait_for_timeout(600)
                ensure_saved_words_loaded(page)
                page_words, _ = extract_words_from_page(page)
                all_words.extend(page_words)
                all_words = unique_preserve_order(all_words)
                if not click_next_if_possible(page):
                    break

            context.close()
            browser.close()
            return all_words, None
    except Exception as e:
        return [], str(e)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    words, err = fetch_saved_words_from_mw(headless=False)
    if err:
        raise SystemExit(f"Error: {err}") from None

    OUT_JSON.write_text(json.dumps(words, indent=2), encoding="utf-8")
    with OUT_CSV.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["word"])
        for word in words:
            w.writerow([word])

    try:
        write_docx(words, OUT_DOCX)
    except Exception as e:
        print(f"WARNING: Failed to write Word doc ({OUT_DOCX}): {e}")

    print(f"Extracted {len(words)} words across paginated pages")
    print(f"Wrote: {OUT_JSON.resolve()}")
    print(f"Wrote: {OUT_CSV.resolve()}")
    print(f"Wrote: {OUT_DOCX.resolve()}")


if __name__ == "__main__":
    main()